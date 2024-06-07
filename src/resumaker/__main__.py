"""CLI for resumaker, a tool for producing resumes from DOCX templates and resume information in
YAML files."""
from pathlib import Path
from typing import Dict, Any
import os
import subprocess

import yaml
from docxcompose.composer import Composer
import docx
import docx.document
import docxtpl

from resumaker import resumaker

RESUMAKER_MODULE_DIR = Path(os.path.dirname(__file__))
RESUMAKER_MODULE_ROOT_DIR = RESUMAKER_MODULE_DIR.parents[1]

WORKING_DIR = Path.cwd().resolve()
OUTPUT_DIR = WORKING_DIR / "output"


class JobArguments:

    def __init__(self, job_path: Path):
        self._job_path = job_path
        self._job_dict = self._load_yaml(self._job_path)
        self.output_root_name: str = self._job_dict["output_root_name"]
        self._resume_body_path: Path = Path(self._job_dict["resume_body"])
        self.resume_body = self._load_yaml(self._resume_body_path.resolve())
        self.sections: list[Path] = [Path(p).resolve() for p in self._job_dict["sections"]]

    def _load_yaml(self, path: Path) -> Dict[str, Any]:
        with path.open("r") as f:
            return yaml.safe_load(f)


def remove_last_paragraph(doc: docx.document.Document):
    """Removes the last paragraph in a docx document"""
    last_p = None
    for el in reversed(doc.element.body):
        if el.tag.endswith("p"):
            last_p = el
            break
    if last_p is None:
        raise RuntimeError("No paragraph found in document")
    last_p.getparent().remove(last_p)


def main(job_path: Path):
    with job_path.open("r") as f:
        job = yaml.safe_load(f)

    body_path = Path(job["resume_body"])
    with body_path.open("r") as f:
        resume_body = yaml.safe_load(f)

    output_path = (OUTPUT_DIR / job["output_root_name"]).with_suffix(".docx")

    # Always inheret the style from the first doc in the sections list.
    # parent_doc = docx.Document(job["sections"].pop())
    parent_doc = docx.Document(job["sections"][0])
    composer = Composer(parent_doc)

    for section_path in job["sections"][1:]:
        remove_last_paragraph(composer.doc)
        print("Adding section:", section_path)
        doc = docx.Document(section_path)
        composer.append(doc)

    # Since `docxtpl` doesn't support creating a `DocxTemplate` from an existing `Document`, we need
    # to save the composed document to a temporary file and then load it back in.
    # TODO: Potentially switch to using `tempfile` for the intermediate file.
    temp_path = OUTPUT_DIR / "temp.docx"
    composer.save(temp_path.as_posix())
    doc = docxtpl.DocxTemplate(temp_path.as_posix())

    resumaker.populate_resume_template(doc, resume_body, output_path)

    output_pdf = output_path.with_suffix(".pdf")
    resumaker.convert_to_pdf(output_path, output_pdf)

    # Open the PDF *in the background* if conversion was successful.
    no_display_pdf = False
    if not no_display_pdf:
        print("Displaying PDF in the background.")
        subprocess.Popen(["evince", output_pdf])

    # Convert to TXT if desired.
    # if convert_to_txt:
    #     resumaker.convert_to_txt(output_docx, output_root.with_suffix(".txt"))


if __name__ == "__main__":
    import argparse
    import datetime

    # parser = argparse.ArgumentParser(
    #     description=("CLI for resumaker, a tool for producing resumes from DOCX templates and "
    #                  "resume information contained in YAML files."))

    # parser.add_argument("yaml_file", type=Path, help="The YAML file containing your resume info.")
    # parser.add_argument("docx_template",
    #                     type=Path,
    #                     help="File path to the resume structural template")
    # parser.add_argument("--output-name",
    #                     "-o",
    #                     type=str,
    #                     help="The root name of the output resume file *without* file extension.",
    #                     default="resume")
    # parser.add_argument("--no-display-pdf", default=False, action="store_true")
    # parser.add_argument("--convert-to-txt", "-c", default=False, action="store_true")

    # args = parser.parse_args()

    # args.yaml_file = args.yaml_file.resolve()
    # args.docx_template = args.docx_template.resolve()

    # # Zero-padded date format such that `ls` maintains chronological order.
    # date_str = datetime.date.today().strftime("%Y%m%d")
    # args.output_name = f"{args.output_name}_{date_str}"

    # main(**vars(args))
    main(
        Path(
            "/home/dcolli23/code/personal_projects/resumaker_personal/jobs/resume_composed_test.yaml"
        ))
